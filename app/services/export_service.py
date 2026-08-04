from io import BytesIO

from docx import Document
from openpyxl import Workbook
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

from ..models import CourseEnrollment, LessonProgress, Submission


def build_student_xlsx(user):
    wb = Workbook()
    ws = wb.active
    ws.title = "Успеваемость"
    ws.append(["Студент", user.full_name])
    ws.append(["Электронная почта", user.email])
    ws.append([])
    ws.append(["Курс", "Прогресс, %", "Статус"])
    for enrollment in user.enrollments:
        ws.append([enrollment.course.title, float(enrollment.progress_percent), enrollment.status])
    ws2 = wb.create_sheet("Решения")
    ws2.append(["Задание", "Статус", "Баллы", "Дата"])
    for submission in Submission.query.filter_by(user_id=user.id).order_by(Submission.submitted_at.desc()).all():
        ws2.append([submission.task.title, submission.status, float(submission.score), submission.submitted_at.isoformat()])
    stream = BytesIO()
    wb.save(stream)
    stream.seek(0)
    return stream


def build_student_docx(user):
    doc = Document()
    doc.add_heading("Индивидуальный отчет студента", level=1)
    doc.add_paragraph(f"Студент: {user.full_name}")
    doc.add_paragraph(f"Электронная почта: {user.email}")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Курс"
    hdr[1].text = "Прогресс"
    hdr[2].text = "Статус"
    for enrollment in user.enrollments:
        cells = table.add_row().cells
        cells[0].text = enrollment.course.title
        cells[1].text = f"{enrollment.progress_percent}%"
        cells[2].text = enrollment.status
    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def build_certificate_pdf(user, course):
    stream = BytesIO()
    page = canvas.Canvas(stream, pagesize=A4)
    width, height = A4
    font_name = "Helvetica"
    bold_font = "Helvetica-Bold"
    for candidate in (
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/Library/Fonts/Arial Unicode.ttf",
    ):
        from pathlib import Path
        if Path(candidate).exists():
            try:
                pdfmetrics.registerFont(TTFont("PythonStudyUnicode", candidate))
                font_name = "PythonStudyUnicode"
                bold_font = "PythonStudyUnicode"
                break
            except Exception:
                pass
    page.setTitle("Сертификат PythonStudy")
    page.setFont(bold_font, 24)
    page.drawCentredString(width / 2, height - 120, "PYTHONSTUDY")
    page.setFont(bold_font, 20)
    page.drawCentredString(width / 2, height - 180, "СЕРТИФИКАТ")
    page.setFont(font_name, 14)
    page.drawCentredString(width / 2, height - 240, user.full_name)
    page.drawCentredString(width / 2, height - 280, "успешно завершил(а) курс")
    page.setFont(bold_font, 16)
    page.drawCentredString(width / 2, height - 320, course.title[:70])
    page.setFont(font_name, 11)
    page.drawCentredString(width / 2, 100, "Сформировано в системе PythonStudy")
    page.showPage()
    page.save()
    stream.seek(0)
    return stream

